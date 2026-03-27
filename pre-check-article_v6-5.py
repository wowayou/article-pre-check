import re
import os
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from pathlib import Path
from dataclasses import dataclass, asdict
from typing import Optional, Tuple, List, Any, Dict, Set
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
import logging
import json
import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

# ==========================================
# 核心数据类 (Dataclass): 规范输出数据结构
# ==========================================
@dataclass
class AuditResult:
    project_folder: str
    file_name_link: str
    domain_status: str
    h1_status: str
    heading_status: str
    missing_images: str
    company_info: str
    tdk_content: str
    tdk_advice: str
    all_links: str
    cleaned_links_count: str
    logs: str

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# ==========================================
# 核心处理引擎
# ==========================================
class SEOSuperEngineV65:
    """企业级 Word 文档 SEO 自动化审查与修复引擎"""
    
    def __init__(self, file_path: Path, target_domain: Optional[str], folder_files: List[str], config: Dict[str, Any]):
        self.file_path = file_path
        self.doc = Document(self.file_path)
        self.folder_files = folder_files
        self.target_domain = self._normalize_domain(target_domain)
        self.config = config  # 接收 GUI 配置面板的参数
        
        # 状态记录器
        self.changes = []
        self.missing_images = []
        self.links_removed_count = 0
        self.external_links_found = 0  # [Bug 3] 记录是否真实存在外链
        self.all_links_found = set()

    def _normalize_domain(self, domain_str: Optional[str]) -> Optional[str]:
        """清洗提取主域名"""
        if not domain_str or str(domain_str).strip() == 'nan': return None
        d = str(domain_str).strip().lower()
        d = re.sub(r'^https?://', '', d)
        d = re.sub(r'^www\.', '', d)
        return d.split('/')[0]

    def _is_external(self, url: str) -> bool:
        """判定是否为非本站外链"""
        if not self.target_domain or not url: return False
        url_lower = url.lower()
        if url_lower.startswith(('#', 'mailto:')): return False
        url_clean = re.sub(r'^https?://', '', url_lower)
        url_clean = re.sub(r'^www\.', '', url_clean)
        return self.target_domain not in url_clean

    def _strip_run_hyperlink_style_robust(self, run_el: Any) -> None:
        """
        [Bug 2] 安全卸载文本节点上的超链接样式 (颜色、下划线)
        Safe stripping of hyperlink styles using local-name()
        """
        rPrs = run_el.xpath('./*[local-name()="rPr"]')
        if rPrs:
            rPr = rPrs[0]
            tags_to_remove = (
                rPr.xpath('./*[local-name()="rStyle"]') + 
                rPr.xpath('./*[local-name()="color"]') + 
                rPr.xpath('./*[local-name()="u"]')
            )
            for tag in tags_to_remove:
                rPr.remove(tag)

    def _clean_links(self, apply_fix: bool) -> None:
        """
        外链清理核心逻辑 (支持标准节点与域代码)
        Clean external links core logic (Supports standard nodes and field codes)
        """
        rels = self.doc.part.rels
        links_to_remove = []
        field_codes_to_remove = []

        # 1. 扫描标准节点与纯文本
        for p in self.doc.paragraphs:
            for hl in p._element.xpath('.//w:hyperlink'):
                rId = hl.get(qn('r:id'))
                if rId in rels and rels[rId]._target:
                    url = rels[rId]._target
                    self.all_links_found.add(url)
                    if self._is_external(url): 
                        links_to_remove.append((hl, url))
                        self.external_links_found += 1
                    
        # 2. 扫描域代码 (Field Codes)
        for instr in self.doc.element.xpath('.//w:instrText'):
            if instr.text and "HYPERLINK" in instr.text:
                url_match = re.search(r'"(https?://[^"]+)"', instr.text)
                if url_match:
                    url = url_match.group(1)
                    self.all_links_found.add(url)
                    if self._is_external(url): 
                        field_codes_to_remove.append((instr, url))
                        self.external_links_found += 1

        # 3. 扫描纯文本 (仅提取，无法直接安全删除)
        full_text = "\n".join(p.text for p in self.doc.paragraphs if p.text)
        for url in re.findall(r'(https?://[^\s<>"]+|www\.[^\s<>"]+)', full_text):
            self.all_links_found.add(url)

        # 仅在配置允许且有目标域名时执行清理
        if self.config['clean_links'] and self.target_domain:
            for hl, url in links_to_remove:
                if apply_fix:
                    try:
                        parent = hl.getparent()
                        if parent is not None:
                            for child in list(hl):
                                if child.tag.endswith('}r'): self._strip_run_hyperlink_style_robust(child)
                                hl.addprevious(child)
                            parent.remove(hl)
                            self.links_removed_count += 1
                    except Exception as e:
                        logging.debug(f"Failed to remove hyperlink: {e}")
                else: self.changes.append(f"待清理标准外链: {url}")
                
            for instr, url in field_codes_to_remove:
                if apply_fix:
                    try:
                        instr.text = instr.text.replace('HYPERLINK', 'QUOTE')
                        parent_r = instr.getparent()
                        if parent_r is not None:
                            curr = parent_r.getnext()
                            while curr is not None:
                                if curr.xpath('.//w:fldChar[@w:fldCharType="end"]'): break
                                if curr.tag.endswith('}r'): self._strip_run_hyperlink_style_robust(curr)
                                curr = curr.getnext()
                        self.links_removed_count += 1
                    except Exception as e:
                        logging.debug(f"Failed to remove field code link: {e}")
                else: self.changes.append(f"待清理域代码外链: {url}")

    def _fix_images_and_headings(self, apply_fix: bool) -> str:
        """处理图片标注与层级跳级"""
        img_re = re.compile(r'(img\.)(.*?)\.(jpg|jpeg|png|bmp|gif|webp)', re.I)
        headings = []
        
        for p in self.doc.paragraphs:
            has_tag = "img." in p.text.lower()
            # 穿透 XML 查找真实图片对象
            has_obj = len(p._element.xpath('.//w:drawing | .//w:pict')) > 0
            
            # 如果图片(标注或对象)被错误地设置为了标题格式
            if (has_tag or has_obj) and p.style.name.startswith(('Heading', '标题')):
                if apply_fix: p.style = 'Normal'
                self.changes.append("样式修复: 图片/标注行降级为正文")
            
            # 处理图片标注文本
            if has_tag:
                matches = list(img_re.finditer(p.text))
                new_text = p.text
                for m in matches:
                    clean_fname = re.sub(r'[\\/:*?"<>|]', '', m.group(2))
                    orig_ext = m.group(3).lower()
                    
                    # 逻辑分流：是否强制 .webp
                    target_ext = "webp" if self.config['force_webp'] else orig_ext
                    fixed_tag = f"{m.group(1)}{clean_fname}.{target_ext}"
                    
                    # 检查文件夹中是否存在该图
                    if f"{clean_fname}.{target_ext}".lower() not in self.folder_files:
                        self.missing_images.append(f"{clean_fname}.{target_ext}")
                        
                    if m.group(0) != fixed_tag:
                        if apply_fix:
                            new_text = new_text.replace(m.group(0), fixed_tag)
                            self.changes.append(f"标注修正: {fixed_tag}")
                if apply_fix: p.text = new_text

            # 收集正常标题用于跳级检查
            if p.style.name.startswith(('Heading', '标题')) and not ("img." in p.text.lower() or has_obj):
                m_level = re.search(r'\d', p.style.name)
                if m_level: headings.append((p, int(m_level.group())))

        h_status = self._fix_heading_hierarchy(headings, apply_fix)
            
        return h_status

    def _fix_heading_hierarchy(self, headings: List[Tuple[Any, int]], apply_fix: bool) -> str:
        """
        [Bug 1 修复] 动态标题修复算法 (基于堆栈/基准追踪)
        """
        h_status = "正常"
        last_fixed_lv = 0
        current_offset = 0

        for p, actual_lv in headings:
            if last_fixed_lv == 0:
                current_offset = 0
                potential_lv = actual_lv
            else:
                potential_lv = actual_lv - current_offset
                
                if potential_lv > last_fixed_lv + 1:
                    current_offset = actual_lv - (last_fixed_lv + 1)
                    potential_lv = actual_lv - current_offset
                elif potential_lv < 1:
                    current_offset = actual_lv - 1 if actual_lv > 1 else 0
                    potential_lv = actual_lv - current_offset
                elif actual_lv <= last_fixed_lv and actual_lv > 1:
                    if actual_lv <= last_fixed_lv:
                        if actual_lv <= last_fixed_lv + 1:
                            current_offset = 0
                            potential_lv = actual_lv
                        else:
                            current_offset = actual_lv - (last_fixed_lv + 1)
                            potential_lv = actual_lv - current_offset

            if potential_lv != actual_lv:
                h_status = "层级跳级修复"
                old_name = p.style.name
                if apply_fix:
                    prefix = "Heading " if "Heading" in old_name else "标题 "
                    try:
                        p.style = f"{prefix}{potential_lv}"
                    except Exception as e:
                        logging.warning(f"Failed to apply style {prefix}{potential_lv}: {e}")
                self.changes.append(f"层级修正: {old_name} -> H{potential_lv}")
            
            last_fixed_lv = potential_lv
        
        return h_status

    def _check_h1_uniqueness(self) -> str:
        """检查 H1 唯一性"""
        h1_count = 0
        for p in self.doc.paragraphs:
            if p.style.name.startswith(('Heading 1', '标题 1')):
                # 排除图片对象的干扰
                if not ("img." in p.text.lower() or len(p._element.xpath('.//w:drawing | .//w:pict')) > 0):
                    h1_count += 1
        
        if h1_count == 0: return "❌ 缺失 H1"
        elif h1_count == 1: return "正常 (1个)"
        else: return f"❌ 异常 ({h1_count}个 H1)"

    def _extract_company_info(self) -> str:
        """提取公司描述信息"""
        for p in self.doc.paragraphs:
            full_txt = "".join(node.text for node in p._element.xpath('.//w:t') if node.text)
            if "co., ltd" in full_txt.lower():
                return p.text.strip()
        return "未发现"

    def process(self, apply_fix=False) -> Tuple[str, str, str, str, str, str]:
        """执行单一文件的全面审计与修复 (充当调度器)"""
        skip_rules_config = self.config.get('skip_rules_config', {})
        file_name_lower = self.file_path.name.lower()
        
        skip_h1_check = False
        skip_links_clean = False
        
        for rule in skip_rules_config.get('rules', []):
            if any(kw.lower() in file_name_lower for kw in rule.get('keywords', [])):
                checks_to_skip = rule.get('skip_checks', [])
                if 'h1_check' in checks_to_skip: skip_h1_check = True
                if 'links_clean' in checks_to_skip: skip_links_clean = True

        if not skip_links_clean:
            self._clean_links(apply_fix)
            
        h_status = self._fix_images_and_headings(apply_fix)
        co_info = self._extract_company_info()
        
        if self.config['enable_seo_check'] and not skip_h1_check:
            h1_status = self._check_h1_uniqueness()
        else:
            h1_status = "未开启检查/触发跳过规则"

        if apply_fix and (self.changes or self.links_removed_count > 0):
            try: self.doc.save(self.file_path)
            except Exception as e: logging.error(f"保存失败 {self.file_path.name}: {e}")
            
        logs = "; ".join(set(self.changes)) if self.changes else "正常"
        all_links_str = "\n".join(list(self.all_links_found))
        miss_imgs_str = ", ".join(set(self.missing_images)) if self.missing_images else "无"
        
        return h1_status, h_status, co_info, all_links_str, miss_imgs_str, logs

# ==========================================
# 工作流与 GUI 管理器
# ==========================================
class SEOWorkflowManagerV65:
    def __init__(self) -> None:
        self.domain_map: Dict[str, str] = {}
        self.results: List[AuditResult] = []
        self.root = tk.Tk()
        self.root.title("SEO Super Engine V6.5 - 控制面板")
        self.root.geometry("480x420")
        self.root.eval('tk::PlaceWindow . center')
        
        # 配置变量
        self.var_clean_links = tk.BooleanVar(value=True)
        self.var_force_webp = tk.BooleanVar(value=False)
        self.var_seo_check = tk.BooleanVar(value=True)
        self.run_status = False

    def _build_gui(self):
        """构建专业级配置界面"""
        frame = ttk.Frame(self.root, padding="20")
        frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(frame, text="⚙️ 核心任务配置", font=("微软雅黑", 12, "bold")).pack(anchor=tk.W, pady=(0, 15))

        cb1 = ttk.Checkbutton(frame, text="开启【自动清理非本站外链】", variable=self.var_clean_links)
        cb1.pack(anchor=tk.W, pady=5)
        ttk.Label(frame, text="  (包含标准链接与隐藏域代码链接)", foreground="gray").pack(anchor=tk.W, pady=(0, 10))

        cb2 = ttk.Checkbutton(frame, text="开启【强制统一图片标注后缀为 .webp】", variable=self.var_force_webp)
        cb2.pack(anchor=tk.W, pady=5)
        ttk.Label(frame, text="  (不勾选则严格校验原后缀如 .jpg 是否与文件夹一致)", foreground="gray").pack(anchor=tk.W, pady=(0, 10))

        cb3 = ttk.Checkbutton(frame, text="开启【深度 SEO 内容审查】", variable=self.var_seo_check)
        cb3.pack(anchor=tk.W, pady=5)
        ttk.Label(frame, text="  (包含 H1 唯一性检查、TDK 字符长度打分评估)", foreground="gray").pack(anchor=tk.W, pady=(0, 15))

        self.var_auto_pack = tk.BooleanVar(value=True)
        cb6 = ttk.Checkbutton(frame, text="开启【修复后自动打包】(等同于一键打包.bat)", variable=self.var_auto_pack)
        cb6.pack(anchor=tk.W, pady=5)
        ttk.Label(frame, text="  (自动将清理后的文档与原图提取至桌面交付文件夹)", foreground="gray").pack(anchor=tk.W, pady=(0, 15))

        btn_frame = ttk.Frame(frame)
        btn_frame.pack(fill=tk.X, pady=20)
        ttk.Button(btn_frame, text="🚀 选择映射表并开始执行", command=self._start_workflow).pack(side=tk.RIGHT, padx=5)
        
        self.root.mainloop()

    def get_tdk_and_validate(self, folder_path: Path, name: str) -> Tuple[str, str]:
        """提取 TDK 内容并根据长度打分建议"""
        try:
            word = name.split()[0].replace('.', '').lower()
            for file_path in folder_path.iterdir():
                if file_path.is_file():
                    fname = file_path.name.lower()
                    if fname.startswith(f"tdk-{word}") or fname == "tdk.docx":
                        doc = Document(file_path)
                        full_txt = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                        
                        advice = "未开启检查"
                        if self.var_seo_check.get():
                            title_len = 0
                            desc_len = 0
                            t_match = re.search(r'(?i)title[:：]\s*(.*)', full_txt)
                            d_match = re.search(r'(?i)description[:：]\s*(.*)', full_txt)
                            
                            advices = []
                            if t_match:
                                title_len = len(t_match.group(1).strip())
                                if not (50 <= title_len <= 60): advices.append(f"Title({title_len}字符)不佳")
                            if d_match:
                                desc_len = len(d_match.group(1).strip())
                                if not (150 <= desc_len <= 160): advices.append(f"Desc({desc_len}字符)不佳")
                            
                            advice = "✅ TDK 长度优良" if not advices else "❌ " + ", ".join(advices)
                            if not t_match and not d_match: advice = "⚠️ 未识别到标准 Title/Desc 标签"
                            
                        return full_txt, advice
        except Exception as e:
            logging.debug(f"Failed to read TDK: {e}")
            return "读取失败", "-"
        return "缺失", "-"

    def _start_workflow(self):
        self.run_status = True
        self.root.withdraw()
        self.root.quit()

    def _load_skip_rules_json(self, root_dir: Path) -> dict:
        """ [Feature 3] 加载目录下的 skip_rules.json """
        rules_path = root_dir / "skip_rules.json"
        if rules_path.exists():
            try:
                with open(rules_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logging.error(f"解析 skip_rules.json 失败: {e}")
        return {}

    def run(self):
        self._build_gui()
        if not self.run_status: return

        # 1. 加载映射表
        path = filedialog.askopenfilename(title="1. 选择【项目域名映射表】Excel", filetypes=[("Excel", "*.xlsx")])
        if not path: return
        try:
            df = pd.read_excel(path)
            self.domain_map = dict(zip(df.iloc[:,0].astype(str).str.strip(), df.iloc[:,1].astype(str).str.strip()))
        except Exception as e:
            messagebox.showerror("错误", f"读取 Excel 失败: {e}")
            return

        # 2. 选择目录
        root_dir_str = filedialog.askdirectory(title="2. 选择文章父文件夹")
        if not root_dir_str: return
        root_dir = Path(root_dir_str)
        
        # 3. 映射健康度校验
        folders = [f.name for f in root_dir.iterdir() if f.is_dir()]
        unmapped_test = [f for f in folders if f not in self.domain_map]
        if self.domain_map and len(unmapped_test) > len(folders) * 0.5:
            msg = f"⚠️ 严重警告：大半文件夹名称未在 Excel 第一列找到映射！\n例如：'{unmapped_test[0]}'\n这将导致外链清理失效。是否继续？"
            if not messagebox.askyesno("映射严重不匹配", msg): return
        
        # [Feature 1] 手动选择跳过的文件
        skipped_files_set = set()
        if getattr(self, 'var_skip_rules', tk.BooleanVar(value=True)).get():
            messagebox.showinfo("跳过规则", "请在接下来的窗口中，选择您想要跳过检查的 Word 文档（可多选）。\n如果不需要跳过，直接点击取消即可。")
            skipped_paths = filedialog.askopenfilenames(
                title="3. 选择要跳过的文档 (可按住 Ctrl/Shift 多选)",
                initialdir=root_dir_str,
                filetypes=[("Word Documents", "*.docx")]
            )
            if skipped_paths:
                skipped_files_set = {str(Path(p).resolve()) for p in skipped_paths}

        config_dict = {
            'clean_links': self.var_clean_links.get(),
            'force_webp': self.var_force_webp.get(),
            'enable_seo_check': self.var_seo_check.get(),
            'skip_rules': getattr(self, 'var_skip_rules', tk.BooleanVar(value=True)).get(),
            'skipped_files': skipped_files_set,
            'skip_rules_config': self._load_skip_rules_json(root_dir),
            'dry_run': getattr(self, 'var_dry_run', tk.BooleanVar(value=False)).get(),
            'auto_pack': getattr(self, 'var_auto_pack', tk.BooleanVar(value=True)).get()
        }

        sandbox_dir = root_dir.parent / f"{root_dir.name}_Cleaned_Output"
        sandbox_dir.mkdir(parents=True, exist_ok=True)
        
        logging.basicConfig(
            filename=sandbox_dir / 'audit.log',
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            force=True
        )
        logging.info(f"Started processing. Sandbox: {sandbox_dir}")

        out_f = sandbox_dir / "SEO审计修复报告_v6.5.xlsx"
        is_dry_run = config_dict['dry_run']

        self.execute_all(root_dir, sandbox_dir, config_dict, apply_fix=False)
        
        if not is_dry_run:
            if messagebox.askyesno("确认修复", f"初次扫描完成！是否执行修复？\n(安全：所有文件将输出至 {sandbox_dir.name})"):
                self.results.clear()
                self.execute_all(root_dir, sandbox_dir, config_dict, apply_fix=True)
                
                if config_dict.get('auto_pack'):
                    self._execute_packing(sandbox_dir, root_dir)
                else:
                    messagebox.showinfo("成功", f"修复完成！请查看沙箱目录:\n{sandbox_dir}")
        else:
            messagebox.showinfo("Dry-run 完成", f"仅审计模式已完成，报告生成于:\n{sandbox_dir}")
            
        pd.DataFrame([asdict(res) for res in self.results]).to_excel(out_f, index=False)
        self.root.destroy()

    def _execute_packing(self, cleaned_dir: Path, original_dir: Path):
        """ [Feature 6] 替代 SEO_Packer_Pro.ps1 的内置打包功能 """
        try:
            timestamp = datetime.datetime.now().strftime("%m%d_%H%M")
            
            messagebox.showinfo("选择打包位置", "请选择最终【交付文件夹】要保存的位置。\n默认会在您选择的目录下创建一个带有时间戳的新文件夹。")
            
            temp_root = tk.Tk()
            temp_root.withdraw()
            temp_root.attributes('-topmost', True)
            
            selected_out_dir = filedialog.askdirectory(
                title="选择打包交付物存放目录",
                initialdir=str(Path(os.path.expanduser("~")) / "Desktop")
            )
            temp_root.destroy()
            
            if not selected_out_dir:
                messagebox.showwarning("已取消打包", "您取消了选择打包目录，自动打包已跳过。修复后的文件仍在沙箱中。")
                return
                
            final_delivery = Path(selected_out_dir) / f"Delivery_{timestamp}"
            exclude_screenshot = messagebox.askyesno("打包选项", "是否排除文件名中包含 'screenshot' 的图片？")
            
            logging.info(f"开始打包交付物，目标路径: {final_delivery}")
            final_delivery.mkdir(parents=True, exist_ok=True)
            
            # 1. 复制清理后的文档
            for root, _, files in os.walk(cleaned_dir):
                for file in files:
                    # FIX: 修复了未闭合字符串，补充了 src 的声明
                    if file.endswith('.docx') and not file.startswith('~'):
                        src = Path(root) / file
                        rel_path = src.relative_to(cleaned_dir)
                        dst = final_delivery / rel_path
                        dst.parent.mkdir(parents=True, exist_ok=True)
                        shutil.copy2(src, dst)
            
            # 2. 复制原始资源(图片)
            valid_exts = {'.webp', '.png', '.jpg', '.jpeg'}
            img_count = 0
            for root, _, files in os.walk(original_dir):
                if cleaned_dir in Path(root).parents or Path(root) == cleaned_dir:
                    continue
                    
                for file in files:
                    ext = Path(file).suffix.lower()
                    if ext in valid_exts:
                        if exclude_screenshot and 'screenshot' in file.lower():
                            continue
                            
                        src = Path(root) / file
                        rel_path = src.relative_to(original_dir)
                        dst = final_delivery / rel_path
                        dst.parent.mkdir(parents=True, exist_ok=True)
                        shutil.copy2(src, dst)
                        img_count += 1
            
            messagebox.showinfo("打包成功", f"修复与打包已完成！\n\n已提取 {img_count} 张图片\n交付文件夹已生成于:\n{final_delivery}")
        except Exception as e:
            logging.error(f"打包过程发生错误: {e}")
            messagebox.showerror("打包失败", f"打包过程发生错误，请查看日志:\n{e}")

    def _process_single_file_task(self, file_path: Path, rel_path: Path, sandbox_dir: Path, domain: str, folder_files: List[str], config: Dict[str, Any], folder_path: Path, project_name: str, apply_fix: bool) -> Optional[AuditResult]:
        """ [Feature 4] 单文件处理任务封装 """
        if config.get('skip_rules') and str(file_path.resolve()) in config.get('skipped_files', set()):
            logging.info(f"Skipped by user selection: {file_path.name}")
            return None

        out_file_path = sandbox_dir / rel_path
        out_file_path.parent.mkdir(parents=True, exist_ok=True)
        
        target_path = file_path
        if apply_fix:
            shutil.copy2(file_path, out_file_path)
            target_path = out_file_path

        try:
            engine = SEOSuperEngineV65(target_path, domain, folder_files, config)
            h1_status, h_status, co_info, links, miss, logs = engine.process(apply_fix=apply_fix)
            tdk_content, tdk_advice = self.get_tdk_and_validate(folder_path, file_path.name)
            
            if getattr(engine, 'external_links_found', 0) > 0:
                cleaned_links_count = str(engine.links_removed_count) if apply_fix else "等待修复"
            else:
                cleaned_links_count = "0"
            
            return AuditResult(
                project_folder=project_name,
                file_name_link=f'=HYPERLINK("{target_path}", "{file_path.name}")',
                domain_status=domain if domain else "❌ 未匹配映射(跳过外链清理)",
                h1_status=h1_status,
                heading_status=h_status,
                missing_images=miss,
                company_info=co_info,
                tdk_content=tdk_content,
                tdk_advice=tdk_advice,
                all_links=links,
                cleaned_links_count=cleaned_links_count,
                logs=logs
            )
        except Exception as e:
            logging.error(f"Error processing {file_path.name}: {e}")
            return None

    def execute_all(self, root_dir: Path, sandbox_dir: Path, config: Dict[str, Any], apply_fix: bool) -> None:
        """ [Bug 4] 使用 os.walk 增强目录扫描; [Feature 4] 使用 ThreadPoolExecutor 并行处理 """
        tasks = []
        with ThreadPoolExecutor(max_workers=os.cpu_count() or 4) as executor:
            for root, _, files in os.walk(root_dir):
                curr_dir = Path(root)
                
                if sandbox_dir in curr_dir.parents or curr_dir == sandbox_dir:
                    continue
                
                try:
                    rel_to_root = curr_dir.relative_to(root_dir)
                    project_name = rel_to_root.parts[0] if rel_to_root.parts else curr_dir.name
                except ValueError:
                    project_name = curr_dir.name
                    
                domain = self.domain_map.get(project_name)
                folder_files = [f.lower() for f in files]
                
                for file_name in files:
                    if not file_name.endswith('.docx') or file_name.startswith(('~', 'TDK')): 
                        continue
                        
                    file_path = curr_dir / file_name
                    rel_path = file_path.relative_to(root_dir)
                    
                    tasks.append(executor.submit(
                        self._process_single_file_task,
                        file_path, rel_path, sandbox_dir, domain, folder_files, config, curr_dir, project_name, apply_fix
                    ))
            
            for future in as_completed(tasks):
                res = future.result()
                if res:
                    self.results.append(res)

if __name__ == "__main__":
    SEOWorkflowManagerV65().run()