import os
import re
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import filedialog, messagebox
import logging
from pathlib import Path

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class SEOSuperEngineV64:
    def __init__(self, file_path, target_domain=None):
        self.file_path = Path(file_path)
        self.doc = Document(self.file_path)
        self.folder_files = [f.name.lower() for f in self.file_path.parent.iterdir() if f.is_file()]
        self.target_domain = self._normalize_domain(target_domain)
        
        self.changes = []
        self.missing_images = []
        self.links_removed_count = 0
        self.all_links_found = set()

    def _normalize_domain(self, domain_str):
        if not domain_str or str(domain_str).strip() == 'nan': return None
        d = str(domain_str).strip().lower()
        d = re.sub(r'^https?://', '', d)
        d = re.sub(r'^www\.', '', d)
        return d.split('/')[0]

    def clean_illegal_chars(self, filename):
        return re.sub(r'[\\/:*?"<>|]', '', filename)

    def is_external(self, url):
        if not self.target_domain or not url: return False
        url_lower = url.lower()
        if url_lower.startswith('#') or url_lower.startswith('mailto:'): return False
        url_clean = re.sub(r'^https?://', '', url_lower)
        url_clean = re.sub(r'^www\.', '', url_clean)
        return self.target_domain not in url_clean

    def _strip_run_hyperlink_style(self, run_el):
        """【V6.4 核心新增】剥离文本节点上的超链接特定样式（颜色、下划线、样式名）"""
        rPr = run_el.find(qn('w:rPr'))
        if rPr is not None:
            # 1. 移除专属 Hyperlink 样式绑定
            rStyle = rPr.find(qn('w:rStyle'))
            if rStyle is not None and rStyle.get(qn('w:val')) == 'Hyperlink':
                rPr.remove(rStyle)
            # 2. 移除强制写入的颜色 (如蓝色)
            color = rPr.find(qn('w:color'))
            if color is not None: 
                rPr.remove(color)
            # 3. 移除强制写入的下划线
            u = rPr.find(qn('w:u'))
            if u is not None: 
                rPr.remove(u)

    def remove_hyperlink(self, hyperlink_el):
        """清理标准链接并刷平格式"""
        try:
            parent = hyperlink_el.getparent()
            if parent is None: return False
            for child in list(hyperlink_el):
                # 如果是文本块(w:r)，清洗其样式
                if child.tag.endswith('}r'):
                    self._strip_run_hyperlink_style(child)
                hyperlink_el.addprevious(child)
            parent.remove(hyperlink_el)
            return True
        except: return False

    def remove_field_code_link(self, instr_el):
        """深度清理隐藏域代码外链并刷平格式"""
        try:
            # 1. 破坏跳转指令
            instr_el.text = instr_el.text.replace('HYPERLINK', 'QUOTE')
            
            # 2. 向下文查找真正的锚文本块并“卸妆”
            parent_r = instr_el.getparent()
            if parent_r is not None:
                curr = parent_r.getnext()
                while curr is not None:
                    # 遇到域代码的闭合标志(end)则停止遍历
                    if curr.xpath('.//w:fldChar[@w:fldCharType="end"]'):
                        break
                    # 如果是文本块，清洗其样式
                    if curr.tag.endswith('}r'):
                        self._strip_run_hyperlink_style(curr)
                    curr = curr.getnext()
            return True
        except: return False

    def extract_links_logic(self):
        rels = self.doc.part.rels
        # 1. 标准节点
        for p in self.doc.paragraphs:
            for hl in p._element.xpath('.//w:hyperlink'):
                rId = hl.get(qn('r:id'))
                if rId in rels and rels[rId]._target:
                    self.all_links_found.add(rels[rId]._target)
        
        # 2. 域代码
        for instr in self.doc.element.xpath('.//w:instrText'):
            if instr.text and "HYPERLINK" in instr.text:
                url_match = re.search(r'"(https?://[^"]+)"', instr.text)
                if url_match: self.all_links_found.add(url_match.group(1))

        # 3. 纯文本
        full_text = "\n".join([p.text for p in self.doc.paragraphs if p.text])
        text_urls = re.findall(r'(https?://[^\s<>"]+|www\.[^\s<>"]+)', full_text)
        for url in text_urls: self.all_links_found.add(url)

    def process(self, apply_fix=False, clean_external=False):
        self.extract_links_logic()

        if clean_external and self.target_domain:
            rels = self.doc.part.rels
            links_to_remove = []
            field_codes_to_remove = []
            
            # A：收集标准链接
            for p_el in self.doc._element.xpath('//w:p'):
                for hl in p_el.xpath('.//w:hyperlink'):
                    rId = hl.get(qn('r:id'))
                    if rId in rels and rels[rId]._target:
                        url = rels[rId]._target
                        if self.is_external(url): links_to_remove.append((hl, url))
                        
            # B：收集域代码链接
            for instr in self.doc.element.xpath('.//w:instrText'):
                if instr.text and "HYPERLINK" in instr.text:
                    url_match = re.search(r'"(https?://[^"]+)"', instr.text)
                    if url_match and self.is_external(url_match.group(1)):
                        field_codes_to_remove.append((instr, url_match.group(1)))

            # C：执行删除并无痕融合
            for hl, url in links_to_remove:
                if apply_fix:
                    if self.remove_hyperlink(hl): self.links_removed_count += 1
                else: self.changes.append(f"待清理标准外链: {url}")
                
            for instr, url in field_codes_to_remove:
                if apply_fix:
                    if self.remove_field_code_link(instr): self.links_removed_count += 1
                else: self.changes.append(f"待清理域代码外链: {url}")

        # 图片与标注处理
        img_re = re.compile(r'(img\.)(.*?)\.(jpg|jpeg|png|bmp|gif|webp)', re.I)
        for p in self.doc.paragraphs:
            has_tag = "img." in p.text.lower()
            has_obj = len(p._element.xpath('.//w:drawing | .//w:pict')) > 0
            if has_tag or has_obj:
                if p.style.name.startswith(('Heading', '标题')):
                    if apply_fix: p.style = 'Normal'
                    self.changes.append("样式修复: 图片行降级为正文")
                if has_tag:
                    matches = list(img_re.finditer(p.text))
                    new_text = p.text
                    for m in matches:
                        clean_fname = self.clean_illegal_chars(m.group(2))
                        fixed_tag = f"{m.group(1)}{clean_fname}.webp"
                        if f"{clean_fname}.webp".lower() not in self.folder_files:
                            self.missing_images.append(f"{clean_fname}.webp")
                        if m.group(0) != fixed_tag:
                            if apply_fix:
                                new_text = new_text.replace(m.group(0), fixed_tag)
                                self.changes.append(f"标注修正: {fixed_tag}")
                    if apply_fix: p.text = new_text

        # 标题层级修复
        headings = []
        for p in self.doc.paragraphs:
            if p.style.name.startswith(('Heading', '标题')) and not "img." in p.text.lower():
                m = re.search(r'\d', p.style.name)
                if m: headings.append((p, int(m.group())))
        
        h_status = "正常"
        if headings:
            last_lv = 0
            for p, curr_lv in headings:
                if last_lv > 0 and curr_lv > last_lv + 1:
                    new_lv = last_lv + 1
                    if apply_fix:
                        pref = "Heading " if "Heading" in p.style.name else "标题 "
                        try: p.style = f"{pref}{new_lv}"
                        except: pass
                    h_status = f"层级跳级(H{curr_lv}->H{new_lv})"
                    self.changes.append(h_status)
                last_lv = curr_lv

        # 公司信息提取
        co_info = "未发现"
        for p in self.doc.paragraphs:
            full_txt = "".join([node.text for node in p._element.xpath('.//w:t') if node.text])
            if "co., ltd" in full_txt.lower():
                co_info = p.text.strip()
                break

        if apply_fix and (self.changes or self.links_removed_count > 0):
            self.doc.save(self.file_path)
            
        return h_status, co_info, "; ".join(set(self.changes)), "\n".join(list(self.all_links_found)), ", ".join(set(self.missing_images))

class SEOWorkflowManagerV64:
    def __init__(self):
        self.domain_map = {}
        self.results = []

    def get_tdk(self, folder_path, name):
        try:
            word = name.split()[0].replace('.', '').lower()
            for file_path in folder_path.iterdir():
                if file_path.is_file():
                    fname = file_path.name.lower()
                    if fname.startswith(f"tdk-{word}") or fname == "tdk.docx":
                        doc = Document(file_path)
                        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except: return "读取失败"
        return "缺失"

    def run(self):
        root = tk.Tk(); root.withdraw()
        
        if messagebox.askyesno("1. 配置", "是否加载【项目域名映射表】Excel？"):
            path = filedialog.askopenfilename(title="选择映射Excel", filetypes=[("Excel", "*.xlsx")])
            if path:
                try:
                    df = pd.read_excel(path)
                    self.domain_map = dict(zip(df.iloc[:,0].astype(str).str.strip(), df.iloc[:,1].astype(str).str.strip()))
                except Exception as e:
                    messagebox.showerror("错误", f"读取 Excel 失败: {e}")
                    return

        root_dir_str = filedialog.askdirectory(title="2. 选择文章父文件夹")
        if not root_dir_str: return
        root_dir = Path(root_dir_str)
        
        folders = [f.name for f in root_dir.iterdir() if f.is_dir()]
        unmapped_test = [f for f in folders if f not in self.domain_map]
        if self.domain_map and len(unmapped_test) > len(folders) * 0.5:
            msg = (f"⚠️ 严重警告：您选择的文件夹中，有大半无法在 Excel 第一列找到对应名称。\n\n"
                   f"例如：文件夹叫 '{unmapped_test[0]}'\n"
                   f"但您的 Excel 中没有这个名字。\n\n"
                   f"这会导致外链清理功能全部失效！请修改 Excel 第一列为您电脑上的【真实文件夹名称】。")
            messagebox.showwarning("映射严重不匹配", msg)
        
        out_f = filedialog.asksaveasfilename(title="3. 保存最终报告", defaultextension=".xlsx", initialfile="SEO全功能审计修复报告_v6.4.xlsx")
        if not out_f: return
        clean_opt = messagebox.askyesno("4. 外链清理", "是否开启【自动删除非本站外链】功能？")

        self.execute_all(root_dir, apply_fix=False, clean_opt=clean_opt)
        
        if messagebox.askyesno("5. 确认修复", "审计已完成，是否执行自动化修复？"):
            self.results = [] 
            self.execute_all(root_dir, apply_fix=True, clean_opt=clean_opt)
            messagebox.showinfo("成功", "修复完成！外链已被无痕抹除，请查看最新导出的 Excel 报告。")
        
        pd.DataFrame(self.results).to_excel(out_f, index=False)

    def execute_all(self, root_dir, apply_fix=False, clean_opt=False):
        for folder_path in root_dir.iterdir():
            if not folder_path.is_dir(): continue
            
            project_name = folder_path.name
            domain = self.domain_map.get(project_name)
            
            for file_path in folder_path.glob("*.docx"):
                if file_path.name.startswith(('~', 'TDK')): continue
                
                engine = SEOSuperEngineV64(file_path, target_domain=domain)
                h_status, co, logs, links, miss = engine.process(apply_fix=apply_fix, clean_external=clean_opt)
                
                self.results.append({
                    "项目(文件夹)": project_name,
                    "文件名": f'=HYPERLINK("{file_path}", "{file_path.name}")',
                    "绑定域名状态": domain if domain else "❌ 名称未匹配(跳过清理)",
                    "文档内所有链接": links,
                    "清理外链数量": engine.links_removed_count if apply_fix else "等待修复",
                    "标题层级": h_status,
                    "缺失图片": miss if miss else "无",
                    "公司描述 (Co., Ltd)": co,
                    "TDK内容": self.get_tdk(folder_path, file_path.name),
                    "操作日志": logs if logs else "正常"
                })

if __name__ == "__main__":
    SEOWorkflowManagerV64().run()